
####
#   USAGE:
#   .final_df is a df of the complete document with table data extracted
#   .compressed_df is the final df but without blank rows
#   .header_text has the text from the header
#   .footer_text has the text from the footer
####

# foundation taken from https://github.com/kmrambo/Python-docx-Reading-paragraphs-tables-and-images-in-document-order-/blob/master/main_code/Para_table_image_extraction.py

###Import all necessary packages
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import *
from docx.text.paragraph import Paragraph
from docx.text.paragraph import Run
import xml.etree.ElementTree as ET
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docxcompose.composer import Composer
from docx import Document as Document_compose
import pandas as pd
from xml.etree import ElementTree
from io import StringIO
import io
import csv
import base64

# for the header extraction use a different module
import docx2python as docx2python
from docx2python.iterators import iter_paragraphs


class Docx_df:
    def __init__(self, filepath):

        # Load the docx file into document object. You can input your own docx file in this step by changing the input path below:
        document = Document(filepath)


        ##This function extracts the tables and paragraphs from the document object
        def iter_block_items(parent):
            """
            Yield each paragraph and table child within *parent*, in document order.
            Each returned value is an instance of either Table or Paragraph. *parent*
            would most commonly be a reference to a main Document object, but
            also works for a _Cell object, which itself can contain paragraphs and tables.
            """
            if isinstance(parent, doctwo):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                raise ValueError("something's not right")

            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)


        # This function extracts the table from the document object as a dataframe
        def read_docx_tables(tab_id=None, **kwargs):
            """
            parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)
            Parameters:
                filename:   file name of a Word Document
                tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                            When [None] - return a list of DataFrames (parse all tables)
                kwargs:     arguments to pass to `pd.read_csv()` function
            Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
            """

            def read_docx_tab(tab, **kwargs):
                vf = io.StringIO()
                writer = csv.writer(vf)
                for row in tab.rows:
                    writer.writerow(cell.text for cell in row.cells)
                vf.seek(0)
                return pd.read_csv(vf, **kwargs)

            #    doc = Document(filename)
            if tab_id is None:
                return [read_docx_tab(tab, **kwargs) for tab in document.tables]
            else:
                try:
                    return read_docx_tab(document.tables[tab_id], **kwargs)
                except IndexError:
                    print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
                    raise


        # The combined_df dataframe will store all the content in document order including images, tables and paragraphs.
        # If the content is an image or a table, it has to be referenced from image_df for images and table_list for tables using the corresponding image or table id that is stored in combined_df
        # And if the content is paragraph, the paragraph text will be stored in combined_df
        combined_df = pd.DataFrame(columns=['para_text', 'table_id', 'style'])
        table_mod = pd.DataFrame(columns=['string_value', 'table_id'])

        # The image_df will consist of base64 encoded image data of all the images in the document
        image_df = pd.DataFrame(columns=['image_index', 'image_rID', 'image_filename', 'image_base64_string'])

        # The table_list is a list consisting of all the tables in the document
        table_list = []
        xml_list = []

        i = 0
        imagecounter = 0

        blockxmlstring = ''
        for block in iter_block_items(document):
            if 'text' in str(block):
                isappend = False

                runboldtext = ''
                for run in block.runs:
                    if run.bold:
                        runboldtext = runboldtext + run.text

                style = str(block.style.name)

                appendtxt = str(block.text)
                appendtxt = appendtxt.replace("\n", "")
                appendtxt = appendtxt.replace("\r", "")
                tabid = 'Novalue'
                paragraph_split = appendtxt.lower().split()

                isappend = True
                for run in block.runs:
                    xmlstr = str(run.element.xml)
                    my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xmlstr), events=['start-ns'])])
                    root = ET.fromstring(xmlstr)
                    # Check if pic is there in the xml of the element. If yes, then extract the image data
                    if 'pic:pic' in xmlstr:
                        xml_list.append(xmlstr)
                        for pic in root.findall('.//pic:pic', my_namespaces):
                            cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
                            name_attr = cNvPr_elem.get("name")
                            blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                            embed_attr = blip_elem.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                            isappend = True
                            appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                            document_part = document.part
                            image_part = document_part.related_parts[embed_attr]
                            image_base64 = base64.b64encode(image_part._blob)
                            image_base64 = image_base64.decode()
                            dftemp = pd.DataFrame(
                                {'image_index': [imagecounter], 'image_rID': [embed_attr], 'image_filename': [name_attr],
                                 'image_base64_string': [image_base64]})
                            image_df = pd.concat([image_df, dftemp])
                            style = 'Novalue'
                        imagecounter = imagecounter + 1

            elif 'table' in str(block):
                isappend = True
                style = 'Novalue'
                appendtxt = str(block)
                tabid = i
                dfs = read_docx_tables(tab_id=i)
                dftemp = pd.DataFrame({'para_text': [appendtxt], 'table_id': [i], 'style': [style]})
                table_mod = pd.concat([table_mod, dftemp])
                table_list.append(dfs)
                i = i + 1
            if isappend:
                dftemp = pd.DataFrame({'para_text': [appendtxt], 'table_id': [tabid], 'style': [style]})
                combined_df = pd.concat([combined_df, dftemp])

        # here we extract the header text
        doc = docx2python.docx2python(filepath)
        self.header_text = ' '.join(iter_paragraphs(doc.header)).strip()
        self.footer_text = ' '.join(iter_paragraphs(doc.footer)).strip()


        self.combined_df = combined_df.reset_index(drop=True)
        self.image_df = image_df.reset_index(drop=True)


        # store the table list
        self.table_list = table_list

        # Now we will merge the table dfs into the main compressed df

        data_merged = {
            'text': [],
            'style': []
        }

        for index, row in self.combined_df.iterrows():
            para_text = row['para_text']
            style = row['style']
            #print(para_text)

            if para_text[0:11] != "<docx.table":
                data_merged['text'].append(para_text)
                data_merged['style'].append(style)
            else:
                # now we get the table
                tab = table_list[row['table_id']]

                #okay so there is a glitch in the table extraction. the first table row is stored as the
                #header. So we will just extract it manually

                header_entry = ", ".join(list(tab.columns))
                data_merged['text'].append(header_entry)
                data_merged['style'].append("table text")

                # then we will extract all of the rows
                for tab_row_index in range(0,len(tab)):
                    tab_row = tab.iloc[tab_row_index, :]
                    tab_row_list = list(tab_row)
                    # adjust for cells that are empty i.e. value nan
                    clean_list = []
                    for cell in tab_row_list:
                        if str(cell) == "nan":
                            cell = ""
                        clean_list.append(str(cell))
                    clean_text = " ".join(clean_list).strip()

                    data_merged['text'].append(clean_text)
                    data_merged['style'].append('table text')

        self.final_df = pd.concat([pd.DataFrame([{'text': self.header_text, 'style': 'Header'}]),
                                   pd.DataFrame(data=data_merged)], ignore_index=True)
        self.df_rows = self.final_df
        # compress the final_df

        self.compressed_df = self.final_df[self.final_df['text'].astype(bool)].copy()
        self.compressed_df.reset_index(drop=True, inplace=True)

if __name__ == "__main__":
    pass
    # fp = project_root + "/docx_files/2012-00202.docx"
    # doc = Docx_df(fp)
    # doc.final_df.to_excel(project_root + "/Spreadsheets/new_docx_analysis/analysis_1.xlsx")
    #
    # print(doc.final_df)
    # print("header text")
    # print(doc.header_text)
    # print("footer text")
    # print(doc.footer_text)
    # print("compressed")
    # print(doc.compressed_df)

